"""
reconstructor.py - Guardado y reconstrucción de documentos corregidos.
Preserva codificación UTF-8 y formatos: DOCX, ODT, RTF, CSV, MD, TXT, HTML.
"""

import csv
from pathlib import Path
from typing import Any
from docx import Document
from logs import logger

class ReconstruccionError(Exception):
    """Excepción de dominio cuando falla el guardado o reconstrucción."""
    pass

def _guardar_docx(texto_markdown: str, ruta_destino: Path) -> None:
    """Reconstruye un documento .docx básico a partir de texto estructurado con inserción de imágenes."""
    import re
    from docx.shared import Inches

    doc = Document()
    lineas = texto_markdown.splitlines()

    for linea in lineas:
        linea_str = linea.strip()
        if not linea_str:
            continue

        # Detección de imágenes Markdown: ![alt](ruta)
        match_img = re.match(r"^!\[(.*?)\]\((.*?)\)$", linea_str)
        if match_img:
            alt_txt = match_img.group(1)
            raw_path = match_img.group(2).strip()
            img_path = Path(raw_path)
            if not img_path.exists():
                img_path = Path(raw_path.replace("/", "\\"))
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path.resolve()), width=Inches(5.5))
                    if alt_txt:
                        p_cap = doc.add_paragraph(alt_txt)
                        p_cap.alignment = 1  # Centrado
                    continue
                except Exception as e_img:
                    logger.warning(f"No se pudo insertar imagen en DOCX: {e_img}")

        if linea_str.startswith("# "):
            doc.add_heading(linea_str[2:].strip(), level=1)
        elif linea_str.startswith("## "):
            doc.add_heading(linea_str[3:].strip(), level=2)
        elif linea_str.startswith("### "):
            doc.add_heading(linea_str[4:].strip(), level=3)
        elif linea_str.startswith("- ") or linea_str.startswith("* "):
            doc.add_paragraph(linea_str[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(linea_str)

    doc.save(str(ruta_destino))

def _guardar_odt(texto_markdown: str, ruta_destino: Path) -> None:
    """Reconstruye un documento OpenDocument Text (.odt) estructurado."""
    from odf.opendocument import OpenDocumentText
    from odf.text import H, P, List, ListItem

    doc: Any = OpenDocumentText()
    doc_text: Any = getattr(doc, "text", doc)
    lineas = texto_markdown.splitlines()
    lista_actual = None

    for linea in lineas:
        linea_str = linea.strip()
        if not linea_str:
            lista_actual = None
            continue

        if linea_str.startswith("# "):
            lista_actual = None
            doc_text.addElement(H(outlinelevel=1, text=linea_str[2:].strip()))
        elif linea_str.startswith("## "):
            lista_actual = None
            doc_text.addElement(H(outlinelevel=2, text=linea_str[3:].strip()))
        elif linea_str.startswith("### "):
            lista_actual = None
            doc_text.addElement(H(outlinelevel=3, text=linea_str[4:].strip()))
        elif linea_str.startswith("- ") or linea_str.startswith("* "):
            if lista_actual is None:
                lista_actual = List()
                doc_text.addElement(lista_actual)
            item = ListItem()
            item.addElement(P(text=linea_str[2:].strip()))
            lista_actual.addElement(item)
        else:
            lista_actual = None
            doc_text.addElement(P(text=linea_str))

    doc.save(str(ruta_destino))

def _guardar_rtf(texto_markdown: str, ruta_destino: Path) -> None:
    """Genera un archivo Rich Text Format (.rtf) válido y estructurado."""
    parrafos_rtf = []

    for linea in texto_markdown.splitlines():
        linea_str = linea.strip()
        if not linea_str:
            parrafos_rtf.append(r"\par")
            continue

        texto_escapado = (
            linea_str.replace("\\", "\\\\")
            .replace("{", "\\{")
            .replace("}", "\\}")
        )
        res = []
        for ch in texto_escapado:
            cp = ord(ch)
            if cp < 128:
                res.append(ch)
            else:
                res.append(f"\\u{cp}?")
        linea_unicode = "".join(res)

        if linea_str.startswith("# "):
            parrafos_rtf.append(r"{\b\fs32 " + linea_unicode[2:].strip() + r"}\par")
        elif linea_str.startswith("## "):
            parrafos_rtf.append(r"{\b\fs28 " + linea_unicode[3:].strip() + r"}\par")
        elif linea_str.startswith("### "):
            parrafos_rtf.append(r"{\b\fs24 " + linea_unicode[4:].strip() + r"}\par")
        elif linea_str.startswith("- ") or linea_str.startswith("* "):
            parrafos_rtf.append(r"{\bullet  " + linea_unicode[2:].strip() + r"}\par")
        else:
            parrafos_rtf.append(linea_unicode + r"\par")

    rtf_completo = (
        r"{\rtf1\ansi\ansicpg1252\deff0\nouicompat"
        r"{\fonttbl{\f0\fnil\fcharset0 Calibri;}}"
        r"\viewkind4\uc1\pard\sa200\sl276\slmult1\f0\fs22 "
        + "\n".join(parrafos_rtf)
        + r"}"
    )

    with open(ruta_destino, "w", encoding="ascii", errors="replace") as f:
        f.write(rtf_completo)

def _guardar_csv(texto_markdown: str, ruta_destino: Path) -> None:
    """Reconstruye un archivo CSV a partir de tablas Markdown."""
    lineas = texto_markdown.splitlines()
    filas_csv = []

    for linea in lineas:
        linea_str = linea.strip()
        if not linea_str or not linea_str.startswith("|"):
            continue
        if "---" in linea_str:
            continue
        celdas = [c.strip() for c in linea_str.strip("|").split("|")]
        filas_csv.append(celdas)

    with open(ruta_destino, "w", encoding="utf-8", newline="") as f:
        escritor = csv.writer(f)
        for fila in filas_csv:
            escritor.writerow(fila)

def _guardar_html(texto_markdown: str, ruta_destino: Path) -> None:
    """Guarda una versión HTML limpia con encoding UTF-8 garantizado."""
    parrafos_html = []
    for p in texto_markdown.split("\n\n"):
        p_str = p.strip()
        if not p_str:
            continue
        if p_str.startswith("# "):
            parrafos_html.append(f"<h1>{p_str[2:]}</h1>")
        elif p_str.startswith("## "):
            parrafos_html.append(f"<h2>{p_str[3:]}</h2>")
        elif p_str.startswith("### "):
            parrafos_html.append(f"<h3>{p_str[4:]}</h3>")
        else:
            import re
            match_img = re.match(r"^!\[(.*?)\]\((.*?)\)$", p_str)
            if match_img:
                alt_txt = html.escape(match_img.group(1))
                src_txt = html.escape(match_img.group(2).strip())
                parrafos_html.append(
                    f'<figure style="text-align: center; margin: 24px 0;">'
                    f'<img src="{src_txt}" alt="{alt_txt}" style="max-width: 100%; height: auto; border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.1);">'
                    f'<figcaption style="font-size: 13px; color: #64748b; margin-top: 6px;">{alt_txt}</figcaption>'
                    f'</figure>'
                )
            else:
                parrafos_html.append(f"<p>{p_str.replace(chr(10), '<br>')}</p>")

    html_completo = (
        "<!DOCTYPE html>\n"
        '<html lang="es">\n'
        "<head>\n"
        '  <meta charset="utf-8">\n'
        "  <title>Documento Corregido</title>\n"
        "  <style>body { font-family: sans-serif; line-height: 1.6; max-width: 800px; margin: 40px auto; padding: 0 20px; }</style>\n"
        "</head>\n"
        "<body>\n"
        + "\n".join(parrafos_html)
        + "\n</body>\n"
        "</html>"
    )
    with open(ruta_destino, "w", encoding="utf-8") as f:
        f.write(html_completo)

def guardar_documento_corregido(
    texto_corregido: str,
    ruta_original: Path,
    ruta_destino: Path
) -> Path:
    """
    Guarda el contenido corregido en la ruta destino manteniendo el formato adecuado.
    
    Args:
        texto_corregido: Texto con corrección de estilo y ortografía.
        ruta_original: Archivo de entrada original.
        ruta_destino: Destino donde se guardará la versión corregida.
        
    Returns:
        Ruta final del archivo generado.
    """
    ruta_destino.parent.mkdir(parents=True, exist_ok=True)
    ext = ruta_original.suffix.lower()

    try:
        if ext == ".docx":
            _guardar_docx(texto_corregido, ruta_destino)
        elif ext == ".odt":
            _guardar_odt(texto_corregido, ruta_destino)
        elif ext == ".rtf":
            _guardar_rtf(texto_corregido, ruta_destino)
        elif ext == ".csv":
            _guardar_csv(texto_corregido, ruta_destino)
        elif ext in {".txt", ".md"}:
            with open(ruta_destino, "w", encoding="utf-8") as f:
                f.write(texto_corregido)
        elif ext == ".html":
            _guardar_html(texto_corregido, ruta_destino)
        else:
            # Para formatos binarios o complejos (.pdf, .pptx, .xlsx, .ppt, .xls, .doc):
            # Guardamos el Markdown corregido preservando el nombre con sufijo .md
            ruta_md = ruta_destino.with_suffix(".corregido.md")
            with open(ruta_md, "w", encoding="utf-8") as f:
                f.write(texto_corregido)
            logger.info(f"Formato complejo {ext} exportado como Markdown corregido a {ruta_md}")
            return ruta_md

        logger.info(f"Documento corregido guardado exitosamente en: {ruta_destino}")
        return ruta_destino

    except Exception as e:
        logger.error(f"Fallo al reconstruir {ruta_destino}: {e}", exc_info=True)
        raise ReconstruccionError(f"No fue posible guardar {ruta_destino}: {e}") from e


def _guardar_pdf(texto_markdown: str, ruta_destino: Path) -> None:
    """Genera un archivo PDF maquetado profesionalmente a partir de Markdown con ReportLab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    import html
    import re

    doc = SimpleDocTemplate(
        str(ruta_destino),
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=12
    )
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#1e293b'),
        spaceBefore=10,
        spaceAfter=6
    )
    h3_style = ParagraphStyle(
        'DocH3',
        parent=styles['Heading3'],
        fontSize=11,
        leading=15,
        textColor=colors.HexColor('#334155'),
        spaceBefore=8,
        spaceAfter=4
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#1e293b'),
        spaceAfter=6
    )
    bullet_style = ParagraphStyle(
        'DocBullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3
    )

    story = []
    for linea in texto_markdown.splitlines():
        l_str = linea.strip()
        if not l_str:
            story.append(Spacer(1, 6))
            continue

        texto_escapado = html.escape(l_str)
        texto_formateado = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', texto_escapado)
        texto_formateado = re.sub(r'\*(.+?)\*', r'<i>\1</i>', texto_formateado)

        # Detección de imágenes Markdown: ![alt](ruta)
        match_img = re.match(r"^!\[(.*?)\]\((.*?)\)$", l_str)
        if match_img:
            alt_txt = match_img.group(1)
            raw_path = match_img.group(2).strip()
            img_path = Path(raw_path)
            if not img_path.exists():
                img_path = Path(raw_path.replace("/", "\\"))
            if img_path.exists():
                try:
                    from PIL import Image as PILImage
                    from reportlab.platypus import Image as RLImage
                    with PILImage.open(img_path) as p_img:
                        orig_w, orig_h = p_img.size
                    max_w = 480.0
                    max_h = 360.0
                    scale = min(max_w / orig_w, max_h / orig_h, 1.0)
                    w_pt = orig_w * scale
                    h_pt = orig_h * scale
                    story.append(RLImage(str(img_path.resolve()), width=w_pt, height=h_pt))
                    if alt_txt:
                        caption_style = ParagraphStyle(
                            'DocCaption',
                            parent=styles['Italic'],
                            fontSize=8,
                            leading=11,
                            textColor=colors.HexColor('#64748b'),
                            alignment=1,
                            spaceAfter=8
                        )
                        story.append(Paragraph(html.escape(alt_txt), caption_style))
                    else:
                        story.append(Spacer(1, 8))
                    continue
                except Exception as e_pdf_img:
                    logger.warning(f"No se pudo insertar imagen en PDF: {e_pdf_img}")

        if l_str.startswith("# "):
            story.append(Paragraph(texto_formateado[2:].strip(), title_style))
        elif l_str.startswith("## "):
            story.append(Paragraph(texto_formateado[3:].strip(), h2_style))
        elif l_str.startswith("### "):
            story.append(Paragraph(texto_formateado[4:].strip(), h3_style))
        elif l_str.startswith("- ") or l_str.startswith("* "):
            story.append(Paragraph(f"&bull; {texto_formateado[2:].strip()}", bullet_style))
        else:
            story.append(Paragraph(texto_formateado, body_style))

    doc.build(story)


def exportar_documento_formato(
    texto_markdown: str,
    ruta_destino: Path,
    formato: str
) -> Path:
    """
    Exporta texto estructurado Markdown al formato específico solicitado
    (.pdf, .docx, .odt, .rtf, .csv, .html, .md, .txt).
    """
    formato_normalizado = formato.lower().strip()
    if not formato_normalizado.startswith("."):
        formato_normalizado = f".{formato_normalizado}"

    ruta_destino_ajustada = ruta_destino.with_suffix(formato_normalizado)
    ruta_destino_ajustada.parent.mkdir(parents=True, exist_ok=True)

    try:
        if formato_normalizado == ".pdf":
            _guardar_pdf(texto_markdown, ruta_destino_ajustada)
        elif formato_normalizado == ".docx":
            _guardar_docx(texto_markdown, ruta_destino_ajustada)
        elif formato_normalizado == ".odt":
            _guardar_odt(texto_markdown, ruta_destino_ajustada)
        elif formato_normalizado == ".rtf":
            _guardar_rtf(texto_markdown, ruta_destino_ajustada)
        elif formato_normalizado == ".csv":
            _guardar_csv(texto_markdown, ruta_destino_ajustada)
        elif formato_normalizado == ".html":
            _guardar_html(texto_markdown, ruta_destino_ajustada)
        else:
            with open(ruta_destino_ajustada, "w", encoding="utf-8") as f:
                f.write(texto_markdown)

        logger.info(f"Documento exportado exitosamente a formato {formato_normalizado}: {ruta_destino_ajustada}")
        return ruta_destino_ajustada
    except Exception as e:
        logger.error(f"Error al exportar a {formato_normalizado}: {e}", exc_info=True)
        raise ReconstruccionError(f"No fue posible exportar a {formato_normalizado}: {e}") from e

