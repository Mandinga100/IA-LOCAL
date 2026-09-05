"""
tests/unit/test_image_surgery.py - Pruebas unitarias para cirugía de imágenes (core/image_surgery.py).
"""

import io
from pathlib import Path
import pytest
from PIL import Image
from docx import Document
from core.image_surgery import (
    listar_imagenes_documento,
    quitar_imagen_docx,
    reemplazar_imagen_docx,
    ImageSurgeryError
)


@pytest.fixture
def docx_con_imagenes(tmp_path: Path) -> Path:
    """Crea un documento DOCX con 2 imágenes reales."""
    ruta_docx = tmp_path / "doc_con_imgs.docx"
    doc = Document()
    doc.add_heading("Reporte con Imágenes", level=1)

    # Crear imagen 1 roja
    img1_path = tmp_path / "img1.png"
    img1 = Image.new("RGB", (120, 80), color="red")
    img1.save(str(img1_path))
    doc.add_paragraph("Primer bloque gráfico:")
    doc.add_picture(str(img1_path))

    # Crear imagen 2 azul
    img2_path = tmp_path / "img2.png"
    img2 = Image.new("RGB", (150, 100), color="blue")
    img2.save(str(img2_path))
    doc.add_paragraph("Segundo bloque gráfico:")
    doc.add_picture(str(img2_path))

    doc.save(str(ruta_docx))
    return ruta_docx


def test_listar_imagenes_documento(docx_con_imagenes: Path, tmp_path: Path):
    dir_assets = tmp_path / "assets_list"
    lista = listar_imagenes_documento(docx_con_imagenes, dir_assets)
    assert len(lista) == 2
    assert lista[0]["posicion"] == 1
    assert lista[0]["formato"].upper() == "PNG"
    assert lista[1]["posicion"] == 2


def test_quitar_imagen_docx(docx_con_imagenes: Path, tmp_path: Path):
    ruta_salida = tmp_path / "doc_sin_img1.docx"
    res = quitar_imagen_docx(docx_con_imagenes, ruta_salida, posicion_1_based=1)
    assert res.exists()

    # Listar en el nuevo docx debe dar 1 imagen
    dir_assets = tmp_path / "assets_check"
    lista = listar_imagenes_documento(res, dir_assets)
    assert len(lista) == 1


def test_reemplazar_imagen_docx(docx_con_imagenes: Path, tmp_path: Path):
    nueva_img = tmp_path / "verde.png"
    img_verde = Image.new("RGB", (200, 200), color="green")
    img_verde.save(str(nueva_img))

    ruta_salida = tmp_path / "doc_reemplazado.docx"
    res = reemplazar_imagen_docx(docx_con_imagenes, ruta_salida, posicion_1_based=1, ruta_nueva_imagen=nueva_img)
    assert res.exists()

    # Comprobar que sigue teniendo 2 imágenes
    dir_assets = tmp_path / "assets_check2"
    lista = listar_imagenes_documento(res, dir_assets)
    assert len(lista) == 2


def test_quitar_imagen_posicion_invalida(docx_con_imagenes: Path, tmp_path: Path):
    ruta_salida = tmp_path / "doc_fail.docx"
    with pytest.raises(ImageSurgeryError):
        quitar_imagen_docx(docx_con_imagenes, ruta_salida, posicion_1_based=99)
