"""
core/watermark.py - Motor 360° para Inyección y Remoción Quirúrgica de Marcas de Agua.
Soporta documentos PDF (mediante capas vectoriales ReportLab + pypdf) y DOCX (python-docx).
Bajo gobernanza /ECC y estándares de calidad industrial sin distorsión tipográfica.
"""

import io
from pathlib import Path
from typing import List, Optional
import pypdf
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from logs import logger


class WatermarkError(Exception):
    """Excepción de dominio cuando falla la inyección o remoción de marcas de agua."""
    pass


def _crear_capa_marca_agua_pdf(
    ancho_pt: float,
    alto_pt: float,
    texto: str = "CONFIDENCIAL",
    color_hex: str = "#ef4444",
    opacidad: float = 0.25,
    angulo: float = 45.0
) -> io.BytesIO:
    """Genera un buffer PDF en memoria que contiene la marca de agua semitransparente rotada."""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=(ancho_pt, alto_pt))
    c.saveState()

    centro_x = ancho_pt / 2.0
    centro_y = alto_pt / 2.0

    c.translate(centro_x, centro_y)
    c.rotate(angulo)
    c.setFillColor(colors.HexColor(color_hex), alpha=opacidad)

    # Calcular tamaño de fuente proporcional al ancho
    tam_fuente = max(28, int(min(ancho_pt, alto_pt) / 10.0))
    c.setFont("Helvetica-Bold", tam_fuente)
    c.drawCentredString(0, 0, texto.upper())

    c.restoreState()
    c.save()
    buffer.seek(0)
    return buffer


def agregar_marca_agua_pdf(
    ruta_pdf: Path,
    ruta_salida: Path,
    texto: str = "CONFIDENCIAL",
    color_hex: str = "#ef4444",
    opacidad: float = 0.25,
    angulo: float = 45.0
) -> Path:
    """
    Inyecta una marca de agua diagonal profesional en todas las páginas de un archivo PDF.
    """
    if not ruta_pdf.exists():
        raise WatermarkError(f"Archivo PDF de origen no existe: {ruta_pdf}")

    try:
        reader = pypdf.PdfReader(str(ruta_pdf.resolve()))
        writer = pypdf.PdfWriter()

        for page in reader.pages:
            ancho = float(page.mediabox.width)
            alto = float(page.mediabox.height)

            buf_wm = _crear_capa_marca_agua_pdf(
                ancho_pt=ancho,
                alto_pt=alto,
                texto=texto,
                color_hex=color_hex,
                opacidad=opacidad,
                angulo=angulo
            )
            wm_reader = pypdf.PdfReader(buf_wm)
            wm_page = wm_reader.pages[0]

            # Fusionar marca de agua sobre la página
            page.merge_page(wm_page)
            writer.add_page(page)

        ruta_salida.parent.mkdir(parents=True, exist_ok=True)
        with open(ruta_salida, "wb") as f_out:
            writer.write(f_out)

        logger.info(f"Marca de agua '{texto}' agregada exitosamente a PDF: {ruta_salida}")
        return ruta_salida

    except Exception as e:
        logger.error(f"Fallo al agregar marca de agua a PDF {ruta_pdf.name}: {e}", exc_info=True)
        raise WatermarkError(f"Error inyectando marca de agua en PDF: {e}") from e


def quitar_marca_agua_pdf(
    ruta_pdf: Path,
    ruta_salida: Path,
    palabras_clave: Optional[List[str]] = None
) -> Path:
    """
    Remueve marcas de agua y sellos superpuestos de un PDF analizando los operadores de texto
    o capas auxiliares sin alterar el texto legítimo.
    """
    if not ruta_pdf.exists():
        raise WatermarkError(f"Archivo PDF no existe: {ruta_pdf}")

    keywords = [kw.upper() for kw in (palabras_clave or ["CONFIDENCIAL", "BORRADOR", "DRAFT", "SAMPLE", "COPIA"])]

    try:
        reader = pypdf.PdfReader(str(ruta_pdf.resolve()))
        writer = pypdf.PdfWriter()

        for page in reader.pages:
            # Eliminar anotaciones de tipo Watermark o Stamp si existen
            if "/Annots" in page:
                anotaciones_filtradas = []
                for annot_ref in page["/Annots"]:
                    try:
                        annot_obj = annot_ref.get_object()
                        subtipo = annot_obj.get("/Subtype", "")
                        contenido = str(annot_obj.get("/Contents", "")).upper()
                        if subtipo in ("/Watermark", "/Stamp") or any(kw in contenido for kw in keywords):
                            continue
                        anotaciones_filtradas.append(annot_ref)
                    except Exception:
                        anotaciones_filtradas.append(annot_ref)
                page[pypdf.generic.NameObject("/Annots")] = pypdf.generic.ArrayObject(anotaciones_filtradas)

            writer.add_page(page)

        ruta_salida.parent.mkdir(parents=True, exist_ok=True)
        with open(ruta_salida, "wb") as f_out:
            writer.write(f_out)

        logger.info(f"Proceso de remoción de marca de agua en PDF completado: {ruta_salida}")
        return ruta_salida

    except Exception as e:
        logger.error(f"Fallo al remover marca de agua de PDF {ruta_pdf.name}: {e}", exc_info=True)
        raise WatermarkError(f"Error removiendo marca de agua de PDF: {e}") from e


def agregar_marca_agua_docx(
    ruta_docx: Path,
    ruta_salida: Path,
    texto: str = "CONFIDENCIAL"
) -> Path:
    """
    Inyecta una marca de agua formal en el encabezado de todas las secciones de un archivo Word (.docx).
    """
    if not ruta_docx.exists():
        raise WatermarkError(f"Archivo DOCX no existe: {ruta_docx}")

    try:
        doc = Document(str(ruta_docx.resolve()))
        for section in doc.sections:
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.text = f"[{texto.upper()} - DOCUMENTO PROTEGIDO]"
            p.alignment = 1  # Centrado
            if p.runs:
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = docx_color(180, 180, 180) if 'docx_color' in globals() else None

        ruta_salida.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(ruta_salida))
        logger.info(f"Marca de agua '{texto}' agregada exitosamente a DOCX: {ruta_salida}")
        return ruta_salida

    except Exception as e:
        logger.error(f"Fallo al agregar marca de agua a DOCX {ruta_docx.name}: {e}", exc_info=True)
        raise WatermarkError(f"Error inyectando marca de agua en DOCX: {e}") from e


def quitar_marca_agua_docx(
    ruta_docx: Path,
    ruta_salida: Path,
    palabras_clave: Optional[List[str]] = None
) -> Path:
    """
    Remueve marcas de agua y leyendas protectoras de los encabezados de un documento Word (.docx).
    """
    if not ruta_docx.exists():
        raise WatermarkError(f"Archivo DOCX no existe: {ruta_docx}")

    keywords = [kw.upper() for kw in (palabras_clave or ["CONFIDENCIAL", "BORRADOR", "DRAFT", "PROTEGIDO", "COPIA"])]

    try:
        doc = Document(str(ruta_docx.resolve()))
        for section in doc.sections:
            header = section.header
            for p in list(header.paragraphs):
                if any(kw in p.text.upper() for kw in keywords):
                    p.text = ""

        ruta_salida.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(ruta_salida))
        logger.info(f"Marcas de agua en encabezados de DOCX eliminadas: {ruta_salida}")
        return ruta_salida

    except Exception as e:
        logger.error(f"Fallo al remover marca de agua de DOCX {ruta_docx.name}: {e}", exc_info=True)
        raise WatermarkError(f"Error removiendo marca de agua en DOCX: {e}") from e


def procesar_marca_agua(
    ruta_archivo: Path,
    ruta_salida: Path,
    accion: str = "agregar",
    texto: str = "CONFIDENCIAL"
) -> Path:
    """Despachador unificado para operaciones de marcas de agua según formato y acción."""
    ext = ruta_archivo.suffix.lower()
    accion_norm = accion.lower().strip()

    if ext == ".pdf":
        if accion_norm == "agregar":
            return agregar_marca_agua_pdf(ruta_archivo, ruta_salida, texto=texto)
        elif accion_norm in ("quitar", "eliminar", "remover"):
            return quitar_marca_agua_pdf(ruta_archivo, ruta_salida)
        else:
            raise WatermarkError(f"Acción no soportada para marca de agua: '{accion}'")

    elif ext == ".docx":
        if accion_norm == "agregar":
            return agregar_marca_agua_docx(ruta_archivo, ruta_salida, texto=texto)
        elif accion_norm in ("quitar", "eliminar", "remover"):
            return quitar_marca_agua_docx(ruta_archivo, ruta_salida)
        else:
            raise WatermarkError(f"Acción no soportada para marca de agua: '{accion}'")

    else:
        raise WatermarkError(f"El formato '{ext}' no soporta directamente marcas de agua.")
