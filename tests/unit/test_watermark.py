"""
tests/unit/test_watermark.py - Pruebas unitarias para el motor de marcas de agua (core/watermark.py).
Verifica inyección y remoción en PDF y DOCX bajo principios de calidad /ECC.
"""

import io
from pathlib import Path
import pytest
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from core.watermark import (
    agregar_marca_agua_pdf,
    quitar_marca_agua_pdf,
    agregar_marca_agua_docx,
    quitar_marca_agua_docx,
    procesar_marca_agua,
    WatermarkError
)


@pytest.fixture
def pdf_dummy(tmp_path: Path) -> Path:
    """Genera un archivo PDF válido para pruebas."""
    ruta = tmp_path / "dummy.pdf"
    c = canvas.Canvas(str(ruta), pagesize=letter)
    c.drawString(100, 700, "Texto legítimo del documento original.")
    c.save()
    return ruta


@pytest.fixture
def docx_dummy(tmp_path: Path) -> Path:
    """Genera un archivo DOCX válido para pruebas."""
    ruta = tmp_path / "dummy.docx"
    doc = Document()
    doc.add_heading("Título Original", level=1)
    doc.add_paragraph("Contenido original del documento de prueba.")
    doc.save(str(ruta))
    return ruta


def test_agregar_marca_agua_pdf_exito(pdf_dummy: Path, tmp_path: Path):
    ruta_salida = tmp_path / "salida_wm.pdf"
    res = agregar_marca_agua_pdf(pdf_dummy, ruta_salida, texto="CONFIDENCIAL")
    assert res.exists()
    assert res.stat().st_size > pdf_dummy.stat().st_size


def test_quitar_marca_agua_pdf_exito(pdf_dummy: Path, tmp_path: Path):
    ruta_salida = tmp_path / "salida_limpia.pdf"
    res = quitar_marca_agua_pdf(pdf_dummy, ruta_salida)
    assert res.exists()
    assert res.stat().st_size > 0


def test_agregar_y_quitar_marca_agua_docx(docx_dummy: Path, tmp_path: Path):
    ruta_wm = tmp_path / "salida_wm.docx"
    res_wm = agregar_marca_agua_docx(docx_dummy, ruta_wm, texto="BORRADOR")
    assert res_wm.exists()

    # Verificar que el header contiene la marca
    doc_wm = Document(str(res_wm))
    header_text = "".join(p.text for p in doc_wm.sections[0].header.paragraphs)
    assert "BORRADOR" in header_text

    # Quitar marca
    ruta_limpia = tmp_path / "salida_limpia.docx"
    res_limpia = quitar_marca_agua_docx(res_wm, ruta_limpia)
    assert res_limpia.exists()
    doc_limpio = Document(str(res_limpia))
    header_limpio = "".join(p.text for p in doc_limpio.sections[0].header.paragraphs)
    assert "BORRADOR" not in header_limpio


def test_procesar_marca_agua_despachador(pdf_dummy: Path, docx_dummy: Path, tmp_path: Path):
    salida_pdf = tmp_path / "proc.pdf"
    res_pdf = procesar_marca_agua(pdf_dummy, salida_pdf, accion="agregar", texto="SECRETO")
    assert res_pdf.exists()

    salida_docx = tmp_path / "proc.docx"
    res_docx = procesar_marca_agua(docx_dummy, salida_docx, accion="agregar", texto="SECRETO")
    assert res_docx.exists()


def test_marca_agua_archivo_inexistente(tmp_path: Path):
    archivo_falso = tmp_path / "no_existe.pdf"
    with pytest.raises(WatermarkError):
        agregar_marca_agua_pdf(archivo_falso, tmp_path / "out.pdf")
