"""
tests/unit/test_format_converter.py - Pruebas unitarias para el transformador universal de formatos.
"""

from pathlib import Path
import pytest
from docx import Document
from core.format_converter import transformar_formato_documento, FormatConverterError


@pytest.fixture
def docx_base(tmp_path: Path) -> Path:
    ruta = tmp_path / "base.docx"
    doc = Document()
    doc.add_heading("Capítulo 1: Fundamentos", level=1)
    doc.add_paragraph("Este es un párrafo de prueba para conversión cruzada.")
    doc.save(str(ruta))
    return ruta


def test_transformar_docx_a_pdf(docx_base: Path, tmp_path: Path):
    salida = tmp_path / "convertido.pdf"
    res = transformar_formato_documento(docx_base, ".pdf", salida)
    assert res.exists()
    assert res.suffix == ".pdf"
    assert res.stat().st_size > 0


def test_transformar_docx_a_html(docx_base: Path, tmp_path: Path):
    salida = tmp_path / "convertido.html"
    res = transformar_formato_documento(docx_base, ".html", salida)
    assert res.exists()
    assert res.suffix == ".html"
    contenido = res.read_text(encoding="utf-8")
    assert "Capítulo 1: Fundamentos" in contenido


def test_transformar_docx_a_md(docx_base: Path, tmp_path: Path):
    salida = tmp_path / "convertido.md"
    res = transformar_formato_documento(docx_base, ".md", salida)
    assert res.exists()
    assert res.suffix == ".md"
    assert "# Capítulo 1: Fundamentos" in res.read_text(encoding="utf-8")


def test_transformar_formato_invalido(docx_base: Path, tmp_path: Path):
    with pytest.raises(FormatConverterError):
        transformar_formato_documento(docx_base, ".xyz")


def test_transformar_archivo_inexistente(tmp_path: Path):
    falso = tmp_path / "no_existe.docx"
    with pytest.raises(FormatConverterError):
        transformar_formato_documento(falso, ".pdf")
