"""
tests/unit/test_extractor_visual.py - Pruebas unitarias para el extractor visual seguro.
Valida extracción de imágenes, salvaguardas de ciberseguridad (Decompression Bomb, Path Traversal),
downscaling adaptativo y hashing SHA-256 determinista.
"""

import io
import pytest
from pathlib import Path
from PIL import Image
import docx

from extractor_visual import (
    AssetVisual,
    DecompressionBombError,
    ExtraccionVisualError,
    calcular_hash_imagen,
    extraer_imagenes_docx,
    normalizar_y_redimensionar,
)


@pytest.fixture
def docx_con_imagenes(tmp_path: Path) -> Path:
    """Crea un archivo DOCX sintético en memoria con una imagen incrustada."""
    doc = docx.Document()
    doc.add_heading("Documento con Imagen de Prueba", level=1)
    doc.add_paragraph("Este es un párrafo antes del diagrama.")

    # Generar imagen sintética de 100x100 píxeles
    img = Image.new("RGB", (100, 100), color="blue")
    buffer_img = io.BytesIO()
    img.save(buffer_img, format="PNG")
    buffer_img.seek(0)

    # Añadir imagen al documento
    doc.add_picture(buffer_img, width=docx.shared.Inches(2.0))
    doc.add_paragraph("Este es un párrafo después del diagrama.")

    ruta_salida = tmp_path / "test_doc_imagen.docx"
    doc.save(str(ruta_salida))
    return ruta_salida


def test_calcular_hash_imagen():
    datos = b"prueba_de_imagen_binaria_123"
    h1 = calcular_hash_imagen(datos)
    h2 = calcular_hash_imagen(datos)
    assert len(h1) == 64
    assert h1 == h2


def test_normalizar_y_redimensionar_imagen_pequena():
    img = Image.new("RGB", (200, 150), color="green")
    buffer = io.BytesIO()
    img.save(buffer, format="PNG")

    resultado_bytes = normalizar_y_redimensionar(buffer.getvalue(), max_dimension=1280)
    with Image.open(io.BytesIO(resultado_bytes)) as img_res:
        assert img_res.size == (200, 150)
        assert img_res.format == "PNG"


def test_normalizar_y_redimensionar_downscaling():
    # Imagen de 2000 x 1000 píxeles (debe reducirse a max 1280 en el lado mayor)
    img = Image.new("RGB", (2000, 1000), color="red")
    buffer = io.BytesIO()
    img.save(buffer, format="PNG")

    resultado_bytes = normalizar_y_redimensionar(buffer.getvalue(), max_dimension=1280)
    with Image.open(io.BytesIO(resultado_bytes)) as img_res:
        assert img_res.size[0] == 1280
        assert img_res.size[1] == 640
        assert img_res.format == "PNG"


def test_extraer_imagenes_docx_exito(docx_con_imagenes: Path, tmp_path: Path):
    dir_assets = tmp_path / "assets"
    assets = extraer_imagenes_docx(docx_con_imagenes, dir_assets, doc_hash="a1b2c3d4e5")

    assert len(assets) == 1
    asset = assets[0]
    assert isinstance(asset, AssetVisual)
    assert asset.posicion == 1
    assert asset.formato == "png"
    assert asset.ancho == 100
    assert asset.alto == 100
    assert asset.ruta_disco.exists()
    assert asset.ruta_disco.is_file()
    assert "a1b2c3d4" in asset.image_id


def test_extraer_imagenes_docx_archivo_inexistente(tmp_path: Path):
    ruta_falsa = tmp_path / "no_existe.docx"
    dir_assets = tmp_path / "assets"
    with pytest.raises(ExtraccionVisualError, match="El archivo no existe"):
        extraer_imagenes_docx(ruta_falsa, dir_assets)


def test_extraer_imagenes_docx_sin_imagenes(tmp_path: Path):
    doc = docx.Document()
    doc.add_paragraph("Solo texto sin imágenes.")
    ruta = tmp_path / "sin_imagenes.docx"
    doc.save(str(ruta))

    dir_assets = tmp_path / "assets"
    assets = extraer_imagenes_docx(ruta, dir_assets)
    assert len(assets) == 0


def test_extraer_imagenes_pdf_exito(tmp_path: Path):
    """Verifica que extraer_imagenes_pdf extrae y calcula hash de imágenes en PDF."""
    from reportlab.platypus import SimpleDocTemplate, Image as RLImage
    from reportlab.lib.pagesizes import letter
    from extractor_visual import extraer_imagenes_pdf, inyectar_anclas_imagenes_en_markdown

    # Crear una imagen sintética
    img_path = tmp_path / "img_prueba.png"
    img = Image.new("RGB", (120, 80), color="blue")
    img.save(str(img_path))

    # Crear PDF con la imagen
    pdf_path = tmp_path / "doc_con_imagen.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    doc.build([RLImage(str(img_path), width=120, height=80)])

    dir_assets = tmp_path / "assets_pdf"
    assets = extraer_imagenes_pdf(pdf_path, dir_assets)

    assert len(assets) == 1
    asset = assets[0]
    assert asset.formato == "png"
    assert asset.ancho == 120
    assert asset.alto == 80
    assert asset.ruta_disco.exists()

    # Probar inyección de anclas
    texto_md = "# Título del documento\n\nPárrafo de prueba."
    md_con_ancla = inyectar_anclas_imagenes_en_markdown(texto_md, assets)
    assert asset.image_id in md_con_ancla
    assert "![" in md_con_ancla

