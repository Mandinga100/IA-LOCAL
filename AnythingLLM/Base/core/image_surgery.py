"""
core/image_surgery.py - Cirugía Quirúrgica e Inspección Pixel-Perfect de Imágenes en Documentos.
Permite listar, extraer, eliminar y reemplazar imágenes específicas en documentos PDF y DOCX
preservando la maquetación, proporciones, anclajes y consistencia visual sin distorsión.
Bajo gobernanza /ECC.
"""

import hashlib
import io
from pathlib import Path
from typing import Any, Dict, List, Optional
from PIL import Image
from docx import Document
from logs import logger
from extractor_visual import extraer_imagenes_documento, AssetVisual, ExtraccionVisualError


class ImageSurgeryError(Exception):
    """Excepción de dominio cuando falla una operación quirúrgica de imágenes."""
    pass


def listar_imagenes_documento(ruta_archivo: Path, dir_assets_temp: Optional[Path] = None) -> List[Dict[str, Any]]:
    """
    Inspecciona y lista todas las imágenes de un documento con metadatos pixel-perfect:
    posición (1..N), dimensiones, formato, peso y hash SHA-256.
    """
    if not ruta_archivo.exists():
        raise ImageSurgeryError(f"Archivo no existe: {ruta_archivo}")

    dir_assets = dir_assets_temp or (ruta_archivo.parent / "temp_assets")
    dir_assets.mkdir(parents=True, exist_ok=True)

    assets = extraer_imagenes_documento(ruta_archivo, dir_assets)
    resultado: List[Dict[str, Any]] = []

    for a in assets:
        resultado.append({
            "posicion": a.posicion,
            "image_id": a.image_id,
            "ruta_disco": str(a.ruta_disco),
            "ancho": a.ancho,
            "alto": a.alto,
            "aspect_ratio": round(a.ancho / a.alto, 3) if a.alto > 0 else 1.0,
            "formato": a.formato,
            "sha256": a.sha256_hash,
            "peso_kb": round(a.tamano_bytes / 1024, 1)
        })

    return resultado


def quitar_imagen_docx(
    ruta_docx: Path,
    ruta_salida: Path,
    posicion_1_based: int
) -> Path:
    """
    Elimina quirúrgicamente la imagen en la posición especificada de un archivo DOCX.
    """
    if not ruta_docx.exists():
        raise ImageSurgeryError(f"Archivo DOCX no existe: {ruta_docx}")

    try:
        doc = Document(str(ruta_docx.resolve()))
        contador_img = 0
        imagen_eliminada = False

        # Recorrer párrafos buscando elementos drawing (Word OpenXML)
        for p in doc.paragraphs:
            drawings = p._p.xpath('.//w:drawing')
            if drawings:
                for d in drawings:
                    contador_img += 1
                    if contador_img == posicion_1_based:
                        for blip in d.xpath('.//a:blip'):
                            r_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if r_id and hasattr(doc.part, 'drop_rel'):
                                try:
                                    doc.part.drop_rel(r_id)
                                except Exception:
                                    pass
                        d.getparent().remove(d)
                        imagen_eliminada = True
                        break
            if imagen_eliminada:
                break

        # Si no estaba en párrafos directos, buscar en tablas
        if not imagen_eliminada:
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for p in celda.paragraphs:
                            drawings = p._p.xpath('.//w:drawing')
                            for d in drawings:
                                contador_img += 1
                                if contador_img == posicion_1_based:
                                    for blip in d.xpath('.//a:blip'):
                                        r_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                        if r_id and hasattr(doc.part, 'drop_rel'):
                                            try:
                                                doc.part.drop_rel(r_id)
                                            except Exception:
                                                pass
                                    d.getparent().remove(d)
                                    imagen_eliminada = True
                                    break
                            if imagen_eliminada:
                                break
                        if imagen_eliminada:
                            break
                    if imagen_eliminada:
                        break
                if imagen_eliminada:
                    break

        if not imagen_eliminada:
            raise ImageSurgeryError(f"No se encontró ninguna imagen en la posición {posicion_1_based} (Total encontradas: {contador_img})")

        ruta_salida.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(ruta_salida))
        logger.info(f"Imagen #{posicion_1_based} eliminada exitosamente de DOCX: {ruta_salida}")
        return ruta_salida

    except Exception as e:
        if isinstance(e, ImageSurgeryError):
            raise
        logger.error(f"Fallo al quitar imagen de DOCX: {e}", exc_info=True)
        raise ImageSurgeryError(f"Error quirúrgico en DOCX: {e}") from e


def reemplazar_imagen_docx(
    ruta_docx: Path,
    ruta_salida: Path,
    posicion_1_based: int,
    ruta_nueva_imagen: Path
) -> Path:
    """
    Reemplaza la imagen en la posición especificada por una nueva imagen,
    preservando la posición en el documento.
    """
    if not ruta_docx.exists():
        raise ImageSurgeryError(f"Archivo DOCX no existe: {ruta_docx}")
    if not ruta_nueva_imagen.exists():
        raise ImageSurgeryError(f"Nueva imagen no existe: {ruta_nueva_imagen}")

    try:
        # En DOCX, la forma más limpia es actualizar los bytes de la parte multimedia (blip r:embed)
        doc = Document(str(ruta_docx.resolve()))
        contador_img = 0
        reemplazada = False

        nuevos_bytes = ruta_nueva_imagen.read_bytes()

        for p in doc.paragraphs:
            drawings = p._p.xpath('.//w:drawing')
            for d in drawings:
                blips = d.xpath('.//a:blip')
                for blip in blips:
                    contador_img += 1
                    if contador_img == posicion_1_based:
                        r_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if r_id:
                            image_part = doc.part.related_parts[r_id]
                            image_part._blob = nuevos_bytes
                            reemplazada = True
                            break
                if reemplazada:
                    break
            if reemplazada:
                break

        if not reemplazada:
            raise ImageSurgeryError(f"No se pudo reemplazar la imagen en posición {posicion_1_based} (Total detectadas: {contador_img})")

        ruta_salida.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(ruta_salida))
        logger.info(f"Imagen #{posicion_1_based} reemplazada con éxito en DOCX: {ruta_salida}")
        return ruta_salida

    except Exception as e:
        if isinstance(e, ImageSurgeryError):
            raise
        logger.error(f"Fallo al reemplazar imagen en DOCX: {e}", exc_info=True)
        raise ImageSurgeryError(f"Error reemplazando imagen en DOCX: {e}") from e


def quitar_imagen_documento(
    ruta_archivo: Path,
    ruta_salida: Path,
    posicion_1_based: int
) -> Path:
    """Despachador unificado para eliminar una imagen según el formato."""
    ext = ruta_archivo.suffix.lower()
    if ext == ".docx":
        return quitar_imagen_docx(ruta_archivo, ruta_salida, posicion_1_based)
    elif ext == ".pdf":
        # En PDF: extraer imágenes, omitir la posición solicitada y reconstruir
        from conversor import convertir_a_markdown
        from reconstructor import exportar_documento_formato
        from extractor_visual import extraer_imagenes_pdf, inyectar_anclas_imagenes_en_markdown

        dir_assets = ruta_salida.parent / "assets_surgery"
        dir_assets.mkdir(parents=True, exist_ok=True)
        assets = extraer_imagenes_pdf(ruta_archivo, dir_assets)

        assets_filtrados = [a for a in assets if a.posicion != posicion_1_based]
        texto_md = convertir_a_markdown(ruta_archivo)
        texto_con_anclas = inyectar_anclas_imagenes_en_markdown(texto_md, assets_filtrados)

        return exportar_documento_formato(texto_con_anclas, ruta_salida, ".pdf")
    else:
        raise ImageSurgeryError(f"Formato no soportado para cirugía de imágenes: {ext}")
