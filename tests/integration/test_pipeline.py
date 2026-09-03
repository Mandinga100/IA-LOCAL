import pytest
import respx
import httpx
from pathlib import Path
from docx import Document
from config import Config
from procesador_lote import ProcesadorLote

@respx.mock
def test_pipeline_e2e_mock(tmp_path: Path):
    dir_entrada = tmp_path / "entrada"
    dir_salida = tmp_path / "salida"
    dir_errores = tmp_path / "errores"
    dir_entrada.mkdir(parents=True)
    
    # 1. Crear documento .txt
    doc_txt = dir_entrada / "notas_reunión.txt"
    doc_txt.write_text("Esta reunion tiene herrores ortograficos y ñandues.", encoding="utf-8")

    # 2. Crear documento .docx básico
    doc_docx = dir_entrada / "informe.docx"
    doc = Document()
    doc.add_heading("Titulo del Informe", level=1)
    doc.add_paragraph("Parrafo con faltas gramaticales.")
    doc.save(str(doc_docx))

    # 3. Crear documento .md
    doc_md = dir_entrada / "readme.md"
    doc_md.write_text("# Documento de prueba\n\nTexto con guion — largo y comillas «latinas».", encoding="utf-8")

    # Mock de Ollama
    respx.post("http://localhost:11434/api/generate").mock(
        return_value=httpx.Response(200, json={"response": "Texto corregido y limpio de errores."})
    )

    config = Config(
        ruta_entrada=dir_entrada,
        ruta_salida=dir_salida,
        ruta_errores=dir_errores,
        ollama_url="http://localhost:11434",
        modelo="qwen2.5:7b"
    )

    procesador = ProcesadorLote(config, tipo_documento="general")
    procesador.ejecutar_lote()

    # Verificaciones
    salida_txt = dir_salida / "notas_reunión.txt"
    salida_docx = dir_salida / "informe.docx"
    salida_md = dir_salida / "readme.md"
    ledger = dir_salida / "historial_procesados.json"

    assert salida_txt.exists()
    assert salida_docx.exists()
    assert salida_md.exists()
    assert ledger.exists()

    # Comprobar contenido UTF-8 de salida
    contenido_txt = salida_txt.read_text(encoding="utf-8")
    assert "Texto corregido y limpio de errores." in contenido_txt

    # Verificar que el ledger registró los 3 archivos
    import json
    with open(ledger, "r", encoding="utf-8") as f:
        datos_ledger = json.load(f)
    assert len(datos_ledger) == 3
