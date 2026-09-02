"""
tests/unit/test_reconstructor.py
Suite de pruebas unitarias para reconstructor.py.
Target de cobertura ECC: ≥85% (reconstructor.py estaba al 51%).
"""

import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock

from reconstructor import (
    guardar_documento_corregido,
    ReconstruccionError,
    _guardar_docx,
    _guardar_html,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
MARKDOWN_SIMPLE = (
    "# Título Principal\n\n"
    "## Sección uno\n\n"
    "Párrafo con texto normal.\n\n"
    "### Sub-sección\n\n"
    "- Viñeta uno\n"
    "- Viñeta dos\n\n"
    "Párrafo final con ñ, tildes y «comillas latinas»."
)


# ---------------------------------------------------------------------------
# Tests de _guardar_docx
# ---------------------------------------------------------------------------
class TestGuardarDocx:
    def test_docx_se_crea(self, tmp_path: Path) -> None:
        """Verifica que _guardar_docx genera un archivo .docx válido."""
        destino = tmp_path / "salida.docx"
        _guardar_docx(MARKDOWN_SIMPLE, destino)
        assert destino.exists()
        assert destino.stat().st_size > 0

    def test_docx_contiene_parrafos(self, tmp_path: Path) -> None:
        """Verifica que el .docx generado tiene párrafos no vacíos."""
        from docx import Document
        destino = tmp_path / "parrafos.docx"
        _guardar_docx(MARKDOWN_SIMPLE, destino)
        doc = Document(str(destino))
        textos = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(textos) >= 3

    def test_docx_texto_vacio_genera_archivo(self, tmp_path: Path) -> None:
        """Un markdown sin párrafos no debe lanzar excepción."""
        destino = tmp_path / "vacio.docx"
        _guardar_docx("", destino)
        assert destino.exists()


# ---------------------------------------------------------------------------
# Tests de _guardar_html
# ---------------------------------------------------------------------------
class TestGuardarHtml:
    def test_html_se_crea(self, tmp_path: Path) -> None:
        """Verifica que _guardar_html genera un archivo .html válido."""
        destino = tmp_path / "salida.html"
        _guardar_html(MARKDOWN_SIMPLE, destino)
        assert destino.exists()

    def test_html_contiene_doctype_y_utf8(self, tmp_path: Path) -> None:
        """Verifica estructura mínima HTML y charset UTF-8."""
        destino = tmp_path / "meta.html"
        _guardar_html(MARKDOWN_SIMPLE, destino)
        contenido = destino.read_text(encoding="utf-8")
        assert "<!DOCTYPE html>" in contenido
        assert 'charset="utf-8"' in contenido

    def test_html_h1_presente(self, tmp_path: Path) -> None:
        """Verifica que el encabezado H1 aparece en el HTML."""
        destino = tmp_path / "h1.html"
        _guardar_html(MARKDOWN_SIMPLE, destino)
        contenido = destino.read_text(encoding="utf-8")
        assert "<h1>" in contenido
        assert "Título Principal" in contenido

    def test_html_preserva_caracteres_especiales(self, tmp_path: Path) -> None:
        """Verifica que ñ, tildes y «» sobreviven la escritura UTF-8."""
        destino = tmp_path / "unicode.html"
        _guardar_html(MARKDOWN_SIMPLE, destino)
        contenido = destino.read_text(encoding="utf-8")
        assert "ñ" in contenido
        assert "«" in contenido


# ---------------------------------------------------------------------------
# Tests de guardar_documento_corregido (función pública)
# ---------------------------------------------------------------------------
class TestGuardarDocumentoCorregido:
    def test_guarda_txt(self, tmp_path: Path) -> None:
        """.txt se guarda directamente con encoding UTF-8."""
        origen = tmp_path / "doc.txt"
        destino = tmp_path / "salida" / "doc.txt"
        origen.write_text("dummy", encoding="utf-8")
        ruta_final = guardar_documento_corregido(
            "Texto corregido con ñ y tildes.", origen, destino
        )
        assert ruta_final == destino
        assert destino.read_text(encoding="utf-8") == "Texto corregido con ñ y tildes."

    def test_guarda_md(self, tmp_path: Path) -> None:
        """.md se guarda directamente con encoding UTF-8."""
        origen = tmp_path / "doc.md"
        destino = tmp_path / "salida" / "doc.md"
        origen.write_text("dummy", encoding="utf-8")
        ruta_final = guardar_documento_corregido("# Título\n\nContenido.", origen, destino)
        assert ruta_final == destino
        assert "Título" in destino.read_text(encoding="utf-8")

    def test_guarda_html(self, tmp_path: Path) -> None:
        """.html delega a _guardar_html."""
        origen = tmp_path / "doc.html"
        destino = tmp_path / "salida" / "doc.html"
        origen.write_text("dummy", encoding="utf-8")
        ruta_final = guardar_documento_corregido(MARKDOWN_SIMPLE, origen, destino)
        assert ruta_final == destino
        assert "<!DOCTYPE html>" in destino.read_text(encoding="utf-8")

    def test_guarda_docx(self, tmp_path: Path) -> None:
        """.docx delega a _guardar_docx."""
        from docx import Document
        origen = tmp_path / "doc.docx"
        doc = Document()
        doc.add_paragraph("Original")
        doc.save(str(origen))
        destino = tmp_path / "salida" / "doc.docx"
        ruta_final = guardar_documento_corregido(MARKDOWN_SIMPLE, origen, destino)
        assert ruta_final == destino
        assert destino.exists()

    def test_formato_binario_exporta_md(self, tmp_path: Path) -> None:
        """Formatos binarios complejos (.pdf, .xlsx) exportan como .corregido.md."""
        origen = tmp_path / "archivo.pdf"
        destino = tmp_path / "salida" / "archivo.pdf"
        origen.write_bytes(b"%PDF-1.4 fake")
        ruta_final = guardar_documento_corregido("Texto corregido.", origen, destino)
        assert ruta_final.suffix == ".md"
        assert "corregido" in ruta_final.name
        assert "Texto corregido." in ruta_final.read_text(encoding="utf-8")

    def test_lanza_reconstruccion_error_en_fallo(self, tmp_path: Path) -> None:
        """Si open() falla, debe lanzarse ReconstruccionError (no silenciar)."""
        origen = tmp_path / "doc.txt"
        origen.write_text("dummy", encoding="utf-8")
        destino = tmp_path / "salida" / "doc.txt"

        with patch("reconstructor.open", side_effect=PermissionError("Acceso denegado")):
            with pytest.raises(ReconstruccionError, match="No fue posible guardar"):
                guardar_documento_corregido("Texto.", origen, destino)

    def test_crea_directorio_destino_si_no_existe(self, tmp_path: Path) -> None:
        """El directorio padre del destino se crea automáticamente."""
        origen = tmp_path / "doc.txt"
        origen.write_text("dummy", encoding="utf-8")
        ruta_anidada = tmp_path / "nivel1" / "nivel2" / "doc.txt"
        guardar_documento_corregido("Texto.", origen, ruta_anidada)
        assert ruta_anidada.exists()
